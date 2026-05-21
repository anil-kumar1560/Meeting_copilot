"""
document_parser.py

Parses PDF, DOCX, and TXT files into plain text strings.
"""

import os


def parse_document(file_path: str, filename: str) -> str:
    """
    Parse a document file and return its text content.

    Args:
        file_path: Absolute path to the file on disk.
        filename:  Original filename (used to determine extension).

    Returns:
        Extracted plain-text content of the document.

    Raises:
        ValueError: If the file extension is not supported.
        RuntimeError: If parsing fails for any reason.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return _parse_pdf(file_path)
    elif ext == ".docx":
        return _parse_docx(file_path)
    elif ext == ".txt":
        return _parse_txt(file_path)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported types: .pdf, .docx, .txt"
        )


def _parse_pdf(file_path: str) -> str:
    """Extract text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        pages_text: list[str] = []
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text("text")
            if text.strip():
                pages_text.append(text)
        doc.close()
        return "\n\n".join(pages_text)
    except Exception as exc:
        raise RuntimeError(f"Failed to parse PDF '{file_path}': {exc}") from exc


def _parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))
        return "\n\n".join(paragraphs)
    except Exception as exc:
        raise RuntimeError(f"Failed to parse DOCX '{file_path}': {exc}") from exc


def _parse_txt(file_path: str) -> str:
    """Read a plain-text file, trying UTF-8 and then latin-1 as fallback."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()
        except Exception as exc:
            raise RuntimeError(
                f"Failed to read TXT file '{file_path}': {exc}"
            ) from exc
    except Exception as exc:
        raise RuntimeError(f"Failed to read TXT file '{file_path}': {exc}") from exc
